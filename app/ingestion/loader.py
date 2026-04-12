"""
Document Loader
Supports PDF, Markdown/Text, DOCX, HTML documents.
Returns LangChain Document objects with rich metadata.

Fixes:
- UTF-8 → Latin-1 → CP-1252 encoding fallback for text files
- DOCX (python-docx) support including table extraction
- Graceful error messages with specific encoding failure details
"""
from __future__ import annotations

import logging
from pathlib import Path
from typing import List

from langchain_core.documents import Document

logger = logging.getLogger(__name__)

# Supported extensions → loader strategy
SUPPORTED_EXTENSIONS = {
    ".pdf":      "pdf",
    ".txt":      "text",
    ".md":       "text",
    ".markdown": "text",
    ".html":     "html",
    ".htm":      "html",
    ".docx":     "docx",
}


# ---------------------------------------------------------------------------
# Individual loaders
# ---------------------------------------------------------------------------

def _load_pdf(path: Path) -> List[Document]:
    from langchain_community.document_loaders import PyPDFLoader
    loader = PyPDFLoader(str(path))
    docs = loader.load()
    for i, doc in enumerate(docs):
        doc.metadata.update({
            "source":      path.name,
            "source_path": str(path),
            "file_type":   "pdf",
            "page":        doc.metadata.get("page", i),
        })
    return docs


def _load_text(path: Path) -> List[Document]:
    """
    Load plain-text / Markdown with automatic encoding detection.

    Tries UTF-8 first (the standard), then falls back to Latin-1 and
    Windows CP-1252 — both common in older medical PDFs converted to text.
    """
    from langchain_community.document_loaders import TextLoader

    encodings = ("utf-8", "latin-1", "cp1252")
    last_exc: Exception | None = None

    for enc in encodings:
        try:
            loader = TextLoader(str(path), encoding=enc)
            docs = loader.load()
            for doc in docs:
                doc.metadata.update({
                    "source":      path.name,
                    "source_path": str(path),
                    "file_type":   path.suffix.lstrip("."),
                    "page":        0,
                })
            return docs
        except UnicodeDecodeError as exc:
            logger.debug(f"Encoding '{enc}' failed for {path.name}: {exc}")
            last_exc = exc
        except Exception as exc:
            logger.error(f"Unexpected error loading {path.name} ({enc}): {exc}")
            raise

    raise ValueError(
        f"Cannot decode '{path.name}' with any supported encoding "
        f"({', '.join(encodings)}). Last error: {last_exc}"
    )


def _load_html(path: Path) -> List[Document]:
    try:
        from langchain_community.document_loaders import BSHTMLLoader
        loader = BSHTMLLoader(str(path))
        docs = loader.load()
    except ImportError:
        logger.warning("beautifulsoup4 not installed; loading HTML as plain text.")
        docs = _load_text(path)
    for doc in docs:
        doc.metadata.update({
            "source":      path.name,
            "source_path": str(path),
            "file_type":   "html",
            "page":        0,
        })
    return docs


def _load_docx(path: Path) -> List[Document]:
    """
    Load a Microsoft Word (.docx) document.

    Extracts body paragraphs and table cells, preserving paragraph structure
    with double-newline separators so the chunker can split on them.
    """
    try:
        from docx import Document as DocxDocument  # python-docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX support. "
            "Install with: pip install python-docx"
        )

    doc_obj = DocxDocument(str(path))
    parts: List[str] = []

    # Body paragraphs
    for para in doc_obj.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # Tables (rendered as pipe-separated rows)
    for table in doc_obj.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    full_text = "\n\n".join(parts)
    if not full_text.strip():
        logger.warning(f"DOCX '{path.name}' appears to be empty.")
        return []

    return [Document(
        page_content=full_text,
        metadata={
            "source":      path.name,
            "source_path": str(path),
            "file_type":   "docx",
            "page":        0,
        },
    )]


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def load_document(path: str | Path) -> List[Document]:
    """Load a single document file. Returns a list of Documents (one per PDF page)."""
    path = Path(path)
    if not path.exists():
        raise FileNotFoundError(f"Document not found: {path}")

    ext = path.suffix.lower()
    strategy = SUPPORTED_EXTENSIONS.get(ext)
    if strategy is None:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS.keys())}"
        )

    logger.info(f"Loading [{strategy.upper()}] {path.name}")
    if strategy == "pdf":
        return _load_pdf(path)
    elif strategy == "text":
        return _load_text(path)
    elif strategy == "html":
        return _load_html(path)
    elif strategy == "docx":
        return _load_docx(path)
    return []


def load_directory(directory: str | Path, recursive: bool = True) -> List[Document]:
    """Batch-load all supported documents from a directory."""
    directory = Path(directory)
    if not directory.is_dir():
        raise NotADirectoryError(f"Not a directory: {directory}")

    pattern = "**/*" if recursive else "*"
    all_docs: List[Document] = []
    errors: list[str] = []

    files = [
        f for f in directory.glob(pattern)
        if f.is_file() and f.suffix.lower() in SUPPORTED_EXTENSIONS
    ]

    logger.info(f"Found {len(files)} documents in {directory}")

    for file in files:
        try:
            docs = load_document(file)
            all_docs.extend(docs)
            logger.info(f"  ✓ {file.name} → {len(docs)} page(s)")
        except Exception as exc:
            logger.error(f"  ✗ {file.name}: {exc}")
            errors.append(f"{file.name}: {exc}")

    if errors:
        logger.warning(f"Failed to load {len(errors)} file(s): {errors}")

    logger.info(f"Total pages/documents loaded: {len(all_docs)}")
    return all_docs
